import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import io
import warnings
import zipfile
import os
from PIL import Image
import tensorflow as tf
from tensorflow.keras.preprocessing.image import ImageDataGenerator
from tensorflow.keras.applications import VGG16, ResNet50, MobileNetV2
from tensorflow.keras.layers import Dense, Flatten, Dropout
from tensorflow.keras.models import Model
from tensorflow.keras.optimizers import Adam
from sklearn.model_selection import train_test_split
from sklearn.metrics import confusion_matrix, classification_report
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# Suppress warnings for cleaner app output
warnings.filterwarnings("ignore")

# --- Session State Initialization ---
# Initialize keys if they don't exist
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False
if "analysis_completed" not in st.session_state:
    st.session_state.analysis_completed = False
if "uploaded_zip_file_key" not in st.session_state:
    st.session_state.uploaded_zip_file_key = 0 # Unique key for file uploader
if "image_data_info" not in st.session_state:
    st.session_state.image_data_info = None # Stores path to extracted dir, class_names, etc.
if "model_trained" not in st.session_state:
    st.session_state.model_trained = False
if "model" not in st.session_state:
    st.session_state.model = None
if "history" not in st.session_state:
    st.session_state.history = None
if "test_generator" not in st.session_state:
    st.session_state.test_generator = None
if "predicted_labels" not in st.session_state:
    st.session_state.predicted_labels = None
if "true_labels" not in st.session_state:
    st.session_state.true_labels = None
if "test_filenames" not in st.session_state:
    st.session_state.test_filenames = None
if "report_generated" not in st.session_state:
    st.session_state.report_generated = False


# --- Helper Functions ---

def clear_session_state_for_new_run():
    """Resets session state variables relevant to a new data upload or analysis."""
    auth_status = st.session_state.get('authenticated', False)
    # Increment the key for the file uploader to clear its state
    current_zip_key = st.session_state.get('uploaded_zip_file_key', 0)

    # Clear all state variables except 'authenticated' and increment file uploader key
    for key in list(st.session_state.keys()):
        if key not in ['authenticated', 'uploaded_zip_file_key']:
            del st.session_state[key]

    st.session_state.authenticated = auth_status
    st.session_state.uploaded_zip_file_key = current_zip_key + 1
    st.session_state.analysis_completed = False
    st.session_state.model_trained = False
    st.session_state.report_generated = False
    st.rerun()

def get_class_names_from_dir(directory):
    """Extracts class names from subfolder names in a given directory."""
    class_names = [d for d in os.listdir(directory) if os.path.isdir(os.path.join(directory, d))]
    class_names.sort() # Ensure consistent order
    return class_names

def extract_zip_file(uploaded_file):
    """Extracts a zip file to a temporary directory and returns its path."""
    temp_dir_name = f"temp_image_data_{uploaded_file.name.replace('.', '_')}"
    extract_path = os.path.join("/tmp", temp_dir_name) # Using /tmp for robustness in cloud environments

    if not os.path.exists(extract_path):
        os.makedirs(extract_path)
    else:
        # Clear existing content if directory already exists
        import shutil
        shutil.rmtree(extract_path)
        os.makedirs(extract_path)

    with zipfile.ZipFile(uploaded_file, 'r') as zip_ref:
        zip_ref.extractall(extract_path)

    # Find the actual data directory inside the zip if there's a single root folder
    # This handles cases where zips contain a single folder like 'dataset/'
    extracted_contents = os.listdir(extract_path)
    if len(extracted_contents) == 1 and os.path.isdir(os.path.join(extract_path, extracted_contents[0])):
        data_dir = os.path.join(extract_path, extracted_contents[0])
    else:
        data_dir = extract_path # Assume images are directly in the root of the zip

    return data_dir

# --- Login Page ---
def login_page():
    st.set_page_config(layout="centered", initial_sidebar_state="auto")
    st.title("Login to Image Classification App")

    # Define valid usernames and passwords
    # IMPORTANT: For deployment, use st.secrets!
    # Example using st.secrets (recommended for production):
    # VALID_CREDENTIALS = {
    #     st.secrets["credentials"]["user1_username"]: st.secrets["credentials"]["user1_password"],
    #     st.secrets["credentials"]["admin_username"]: st.secrets["credentials"]["admin_password"]
    # }
    # For local testing without secrets.toml (for demonstration):
    VALID_CREDENTIALS = {
        "user1": "pass123",
        "admin": "secure_admin_pass"
    }

    with st.form("login_form"):
        st.subheader("Enter your credentials")
        username = st.text_input("Username")
        password = st.text_input("Password", type="password")
        login_button = st.form_submit_button("Login")

        if login_button:
            if username in VALID_CREDENTIALS and VALID_CREDENTIALS[username] == password:
                st.session_state.authenticated = True
                st.success("Logged in successfully!")
                st.rerun()
            else:
                st.error("Invalid username or password")

# --- Main Application Logic ---
def main_app():
    st.set_page_config(layout="wide", initial_sidebar_state="expanded")
    
    # Sidebar
    with st.sidebar:
        try:
            st.image("SsoLogo.jpg", use_column_width=True) # Ensure SsoLogo.jpg is in your repo
        except FileNotFoundError:
            st.warning("SsoLogo.jpg not found. Please ensure it's in the root directory.")
        st.title("Image Classification App")
        st.markdown("---")

        if st.button("🗑️ Clear All Data & Start Fresh", help="Clears all uploaded data and trained models."):
            clear_session_state_for_new_run()

        if st.button("Logout"):
            st.session_state.authenticated = False
            st.info("Logged out successfully.")
            st.rerun()

    # Main Content
    st.header("Upload Image Dataset (ZIP File)")
    st.info("""
        Please upload a `.zip` file containing your image dataset.
        **Important:** Organize your images into subfolders, where each subfolder name represents a class label.
        Example structure:
        ```
        your_dataset.zip/
        ├── class_A/
        │   ├── image1.jpg
        │   └── image2.png
        └── class_B/
            ├── image3.jpeg
            └── image4.jpg
        ```
    """)

    uploaded_file = st.file_uploader(
        "Choose a ZIP file...",
        type="zip",
        key=f"zip_uploader_{st.session_state.uploaded_zip_file_key}"
    )

    if uploaded_file is not None and st.session_state.image_data_info is None:
        with st.spinner("Extracting and analyzing dataset..."):
            try:
                # Store the uploaded file in session state to avoid re-uploading on rerun
                st.session_state.uploaded_file_obj = uploaded_file
                extracted_data_dir = extract_zip_file(uploaded_file)
                class_names = get_class_names_from_dir(extracted_data_dir)

                if not class_names:
                    st.error("No class subfolders found in the ZIP file. Please ensure your images are organized into subfolders, with each subfolder representing a class.")
                    st.session_state.uploaded_file_obj = None # Clear to allow re-upload
                    return

                total_images = sum([len(files) for r, d, files in os.walk(extracted_data_dir)])

                st.session_state.image_data_info = {
                    "extracted_path": extracted_data_dir,
                    "class_names": class_names,
                    "num_classes": len(class_names),
                    "total_images": total_images
                }
                st.success("Dataset extracted and analyzed successfully!")
                st.rerun() # Rerun to update the UI with extracted info
            except zipfile.BadZipFile:
                st.error("Invalid ZIP file. Please upload a valid .zip archive.")
                st.session_state.uploaded_file_obj = None
            except Exception as e:
                st.error(f"An error occurred during extraction: {e}")
                st.session_state.uploaded_file_obj = None

    if st.session_state.image_data_info:
        info = st.session_state.image_data_info
        st.subheader("Dataset Overview")
        st.write(f"**Extracted Path:** `{info['extracted_path']}`")
        st.write(f"**Number of Classes:** {info['num_classes']}")
        st.write(f"**Class Names:** {', '.join(info['class_names'])}")
        st.write(f"**Total Images Found:** {info['total_images']}")

        st.markdown("---")
        st.subheader("Data Preprocessing & Augmentation")
        col_img_size, col_split = st.columns(2)
        with col_img_size:
            image_size = st.selectbox(
                "Select Image Size (Width x Height):",
                options=[(64, 64), (128, 128), (224, 224)],
                format_func=lambda x: f"{x[0]}x{x[1]}",
                index=2 # Default to 224x224 for better performance with pre-trained models
            )
        with col_split:
            test_size = st.slider(
                "Test Set Split Ratio:",
                min_value=0.1, max_value=0.5, value=0.2, step=0.05,
                help="Percentage of data to reserve for testing the model."
            )

        st.markdown("---")
        st.subheader("Data Augmentation (Optional)")
        st.write("Data augmentation helps improve model generalization by creating new training samples from existing ones.")
        col_aug1, col_aug2, col_aug3 = st.columns(3)
        with col_aug1:
            rotation_range = st.slider("Rotation Range (degrees):", 0, 90, 0, help="Random rotation in degrees.")
            width_shift_range = st.slider("Width Shift Range (0.0-1.0):", 0.0, 0.5, 0.0, step=0.05, help="Fraction of total width.")
        with col_aug2:
            height_shift_range = st.slider("Height Shift Range (0.0-1.0):", 0.0, 0.5, 0.0, step=0.05, help="Fraction of total height.")
            shear_range = st.slider("Shear Range (degrees):", 0.0, 0.5, 0.0, step=0.05, help="Shear intensity (shear angle in radians).")
        with col_aug3:
            zoom_range = st.slider("Zoom Range (0.0-1.0):", 0.0, 0.5, 0.0, step=0.05, help="Range for random zoom.")
            horizontal_flip = st.checkbox("Horizontal Flip", value=True, help="Randomly flip inputs horizontally.")
            
        st.markdown("---")
        st.subheader("Model Configuration & Training")
        col_model, col_epochs, col_batch_size = st.columns(3)
        with col_model:
            model_architecture = st.selectbox(
                "Select CNN Architecture:",
                options=["VGG16", "ResNet50", "MobileNetV2"],
                index=1, # Default to ResNet50
                help="Pre-trained models offer excellent performance, especially with limited data."
            )
        with col_epochs:
            epochs = st.number_input("Epochs:", min_value=5, max_value=50, value=10, step=5,
                                     help="Number of times the model will iterate over the entire dataset.")
        with col_batch_size:
            batch_size = st.number_input("Batch Size:", min_value=16, max_value=128, value=32, step=16,
                                          help="Number of samples per gradient update.")

        learning_rate = st.number_input("Learning Rate (Adam Optimizer):", min_value=0.00001, max_value=0.01, value=0.001, format="%f",
                                        help="Controls how much the model's weights are adjusted with respect to the loss gradient.")


        if st.button("🚀 Train Image Classifier", key="train_button", disabled=st.session_state.model_trained):
            if info['num_classes'] < 2:
                st.error("Your dataset must contain at least two classes (subfolders) to perform classification.")
                st.stop()

            with st.spinner("Preparing data and training model... This may take a while."):
                # --- Data Generators ---
                datagen = ImageDataGenerator(
                    rescale=1./255, # Normalize pixel values to [0, 1]
                    rotation_range=rotation_range,
                    width_shift_range=width_shift_range,
                    height_shift_range=height_shift_range,
                    shear_range=shear_range,
                    zoom_range=zoom_range,
                    horizontal_flip=horizontal_flip,
                    validation_split=test_size # Use test_size for validation split
                )

                train_generator = datagen.flow_from_directory(
                    info['extracted_path'],
                    target_size=image_size,
                    batch_size=batch_size,
                    class_mode='categorical', # For multi-class classification
                    subset='training',
                    seed=42
                )

                test_generator = datagen.flow_from_directory(
                    info['extracted_path'],
                    target_size=image_size,
                    batch_size=batch_size,
                    class_mode='categorical',
                    subset='validation', # Using the 'validation' split as our test set
                    seed=42,
                    shuffle=False # Important for consistent evaluation order
                )

                # Store test_generator for later evaluation
                st.session_state.test_generator = test_generator

                # --- Build Model (Transfer Learning) ---
                if model_architecture == "VGG16":
                    base_model = VGG16(weights='imagenet', include_top=False, input_shape=(image_size[0], image_size[1], 3))
                elif model_architecture == "ResNet50":
                    base_model = ResNet50(weights='imagenet', include_top=False, input_shape=(image_size[0], image_size[1], 3))
                elif model_architecture == "MobileNetV2":
                    base_model = MobileNetV2(weights='imagenet', include_top=False, input_shape=(image_size[0], image_size[1], 3))

                # Freeze the base model layers
                base_model.trainable = False

                # Add custom classification head
                x = base_model.output
                x = Flatten()(x)
                x = Dense(256, activation='relu')(x)
                x = Dropout(0.5)(x)
                predictions = Dense(info['num_classes'], activation='softmax')(x) # Output layer with softmax

                model = Model(inputs=base_model.input, outputs=predictions)

                model.compile(optimizer=Adam(learning_rate=learning_rate),
                              loss='categorical_crossentropy',
                              metrics=['accuracy'])

                # --- Train Model ---
                history = model.fit(
                    train_generator,
                    epochs=epochs,
                    validation_data=test_generator,
                    verbose=0 # Suppress verbose output in Streamlit
                )

                st.session_state.model = model
                st.session_state.history = history
                st.session_state.model_trained = True
                st.session_state.analysis_completed = True
                st.success("Model training complete!")
                st.rerun()

    if st.session_state.model_trained:
        st.markdown("---")
        st.header("📊 Model Evaluation")

        model = st.session_state.model
        history = st.session_state.history
        test_generator = st.session_state.test_generator
        class_names = st.session_state.image_data_info['class_names']

        # Get true labels and predicted labels
        y_true = test_generator.classes
        y_pred_probs = model.predict(test_generator, verbose=0)
        y_pred_labels = np.argmax(y_pred_probs, axis=1)

        # Ensure that the predicted and true labels align with the generator's internal mapping
        # test_generator.class_indices maps class names to integer labels.
        # We need to ensure class_names list matches this order for report
        true_class_labels = [class_names[label] for label in y_true]
        predicted_class_labels = [class_names[label] for label in y_pred_labels]

        st.session_state.predicted_labels = predicted_class_labels
        st.session_state.true_labels = true_class_labels
        st.session_state.test_filenames = test_generator.filenames


        st.subheader("Training History (Loss & Accuracy)")
        fig, ax = plt.subplots(1, 2, figsize=(15, 5))

        ax[0].plot(history.history['loss'], label='Training Loss')
        ax[0].plot(history.history['val_loss'], label='Validation Loss')
        ax[0].set_title('Loss over Epochs')
        ax[0].set_xlabel('Epoch')
        ax[0].set_ylabel('Loss')
        ax[0].legend()

        ax[1].plot(history.history['accuracy'], label='Training Accuracy')
        ax[1].plot(history.history['val_accuracy'], label='Validation Accuracy')
        ax[1].set_title('Accuracy over Epochs')
        ax[1].set_xlabel('Epoch')
        ax[1].set_ylabel('Accuracy')
        ax[1].legend()

        st.pyplot(fig)
        plt.close(fig) # Close plot to free memory

        st.subheader("Classification Report")
        report = classification_report(y_true, y_pred_labels, target_names=class_names, output_dict=False)
        st.text(report)

        st.subheader("Confusion Matrix")
        cm = confusion_matrix(y_true, y_pred_labels)
        fig_cm, ax_cm = plt.subplots(figsize=(8, 6))
        sns.heatmap(cm, annot=True, fmt='d', cmap='Blues', xticklabels=class_names, yticklabels=class_names, ax=ax_cm)
        ax_cm.set_xlabel('Predicted')
        ax_cm.set_ylabel('True')
        ax_cm.set_title('Confusion Matrix')
        st.pyplot(fig_cm)
        plt.close(fig_cm)

        st.subheader("Sample Predictions from Test Set")
        num_samples_to_show = min(9, len(test_generator.filenames)) # Show up to 9 samples
        if num_samples_to_show > 0:
            fig_samples, axes = plt.subplots(3, 3, figsize=(10, 10))
            axes = axes.flatten()

            for i in range(num_samples_to_show):
                # Get original image path from test_generator
                original_img_path = os.path.join(test_generator.directory, test_generator.filenames[i])
                img = Image.open(original_img_path).resize(image_size) # Resize for display consistency

                true_label_name = class_names[test_generator.classes[i]]
                predicted_label_name = class_names[np.argmax(y_pred_probs[i])]
                confidence = np.max(y_pred_probs[i]) * 100

                color = "green" if true_label_name == predicted_label_name else "red"

                axes[i].imshow(img)
                axes[i].set_title(f"True: {true_label_name}\nPred: {predicted_label_name}\nConf: {confidence:.2f}%", color=color, fontsize=10)
                axes[i].axis('off')
            
            # Hide unused subplots
            for j in range(num_samples_to_show, 9):
                fig_samples.delaxes(axes[j])

            plt.tight_layout()
            st.pyplot(fig_samples)
            plt.close(fig_samples)
        else:
            st.info("No images available in the test set to display samples.")

        st.markdown("---")
        st.header("🔮 Make a New Prediction ('What is this image?')")
        
        new_image_file = st.file_uploader("Upload a single image for prediction:", type=["jpg", "jpeg", "png"])

        if new_image_file is not None:
            with st.spinner("Predicting..."):
                try:
                    img = Image.open(new_image_file).convert('RGB')
                    
                    # Preprocess for model input
                    img_resized = img.resize(image_size) # Use the same size as training
                    img_array = np.array(img_resized) / 255.0 # Normalize pixel values
                    img_array = np.expand_dims(img_array, axis=0) # Add batch dimension

                    prediction = model.predict(img_array)
                    predicted_class_index = np.argmax(prediction)
                    predicted_class_name = class_names[predicted_class_index]
                    confidence_score = prediction[0][predicted_class_index] * 100

                    st.image(img, caption='Uploaded Image', use_column_width=False, width=250)
                    st.success(f"**Predicted Class:** {predicted_class_name}")
                    st.write(f"**Confidence:** {confidence_score:.2f}%")

                    st.subheader("All Class Probabilities:")
                    prob_df = pd.DataFrame({
                        'Class': class_names,
                        'Probability': prediction[0]
                    }).sort_values(by='Probability', ascending=False).reset_index(drop=True)
                    prob_df['Probability'] = prob_df['Probability'].map(lambda x: f"{x:.4f}")
                    st.dataframe(prob_df, hide_index=True, width=300)

                except Exception as e:
                    st.error(f"Error making prediction: {e}")

        st.markdown("---")
        st.header("📝 Generate Report")
        if st.button("Generate Classification Report (Word Document)", disabled=st.session_state.report_generated):
            with st.spinner("Generating report..."):
                doc = Document()
                doc.add_heading('Image Classification Analysis Report', level=1)

                # Header for SSO Consultants (You can customize this)
                try:
                    doc.add_picture('SsoLogo.jpg', width=Inches(1.0))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except FileNotFoundError:
                    doc.add_paragraph("SSO Consultants") # Fallback if logo not found

                doc.add_paragraph(f"Report Generated: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M:%S')}")
                doc.add_page_break()

                doc.add_heading('1. Dataset Overview', level=2)
                doc.add_paragraph(f"Number of Classes: {info['num_classes']}")
                doc.add_paragraph(f"Class Names: {', '.join(info['class_names'])}")
                doc.add_paragraph(f"Total Images: {info['total_images']}")
                doc.add_paragraph(f"Image Size used for training: {image_size[0]}x{image_size[1]}")
                doc.add_paragraph(f"Test Set Split Ratio: {test_size*100:.0f}%")
                doc.add_page_break()

                doc.add_heading('2. Data Preprocessing & Augmentation', level=2)
                doc.add_paragraph("Images were resized and normalized. The following augmentation techniques were applied:")
                aug_details = [
                    f"- Rotation Range: {rotation_range} degrees",
                    f"- Width Shift Range: {width_shift_range}",
                    f"- Height Shift Range: {height_shift_range}",
                    f"- Shear Range: {shear_range}",
                    f"- Zoom Range: {zoom_range}",
                    f"- Horizontal Flip: {'Enabled' if horizontal_flip else 'Disabled'}"
                ]
                for detail in aug_details:
                    doc.add_paragraph(detail, style='List Bullet')
                doc.add_page_break()

                doc.add_heading('3. Model Configuration', level=2)
                doc.add_paragraph(f"Model Architecture: {model_architecture} (Transfer Learning)")
                doc.add_paragraph(f"Epochs: {epochs}")
                doc.add_paragraph(f"Batch Size: {batch_size}")
                doc.add_paragraph(f"Learning Rate: {learning_rate}")
                doc.add_paragraph("Optimizer: Adam")
                doc.add_page_break()

                doc.add_heading('4. Model Evaluation Results', level=2)

                # Training History Plots
                doc.add_heading('4.1 Training History (Loss & Accuracy)', level=3)
                fig_hist, ax_hist = plt.subplots(1, 2, figsize=(12, 5))
                ax_hist[0].plot(history.history['loss'], label='Training Loss')
                ax_hist[0].plot(history.history['val_loss'], label='Validation Loss')
                ax_hist[0].set_title('Loss over Epochs')
                ax_hist[0].set_xlabel('Epoch')
                ax_hist[0].set_ylabel('Loss')
                ax_hist[0].legend()
                ax_hist[1].plot(history.history['accuracy'], label='Training Accuracy')
                ax_hist[1].plot(history.history['val_accuracy'], label='Validation Accuracy')
                ax_hist[1].set_title('Accuracy over Epochs')
                ax_hist[1].set_xlabel('Epoch')
                ax_hist[1].set_ylabel('Accuracy')
                ax_hist[1].legend()
                plt.tight_layout()
                img_buf = io.BytesIO()
                fig_hist.savefig(img_buf, format='png', bbox_inches='tight')
                img_buf.seek(0)
                doc.add_picture(img_buf, width=Inches(6.0))
                plt.close(fig_hist) # Close plot

                # Classification Report
                doc.add_heading('4.2 Classification Report', level=3)
                report_text = classification_report(st.session_state.true_labels, st.session_state.predicted_labels, output_dict=False)
                doc.add_paragraph(report_text)

                # Confusion Matrix
                doc.add_heading('4.3 Confusion Matrix', level=3)
                cm = confusion_matrix(st.session_state.true_labels, st.session_state.predicted_labels, labels=class_names)
                fig_cm_rep, ax_cm_rep = plt.subplots(figsize=(8, 6))
                sns.heatmap(cm, annot=True, fmt='d', cmap='Blues', xticklabels=class_names, yticklabels=class_names, ax=ax_cm_rep)
                ax_cm_rep.set_xlabel('Predicted')
                ax_cm_rep.set_ylabel('True')
                ax_cm_rep.set_title('Confusion Matrix')
                img_buf_cm = io.BytesIO()
                fig_cm_rep.savefig(img_buf_cm, format='png', bbox_inches='tight')
                img_buf_cm.seek(0)
                doc.add_picture(img_buf_cm, width=Inches(5.0))
                plt.close(fig_cm_rep) # Close plot
                doc.add_page_break()

                doc.add_heading('4.4 Sample Predictions', level=3)
                doc.add_paragraph("A selection of images from the test set with their true and predicted labels:")
                num_samples_report = min(6, len(st.session_state.test_filenames)) # Show up to 6 samples in report
                if num_samples_report > 0:
                    for i in range(num_samples_report):
                        original_img_path = os.path.join(test_generator.directory, st.session_state.test_filenames[i])
                        img = Image.open(original_img_path).convert('RGB')

                        true_label = st.session_state.true_labels[i]
                        predicted_label = st.session_state.predicted_labels[i]

                        # Display image and details
                        img_buf_sample = io.BytesIO()
                        img.save(img_buf_sample, format='PNG')
                        img_buf_sample.seek(0)
                        doc.add_picture(img_buf_sample, width=Inches(1.5))
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.add_run(f"   True: {true_label}, Predicted: {predicted_label}")
                        doc.add_paragraph("") # Add a blank line for spacing
                else:
                    doc.add_paragraph("No sample images available from the test set.")
                
                doc.add_page_break()

                doc.add_heading('5. Conclusion and Recommendations', level=2)
                doc.add_paragraph("The model achieved satisfactory performance for image classification. Further improvements could involve:")
                doc.add_paragraph("- Collecting more diverse training data.")
                doc.add_paragraph("- Fine-tuning the pre-trained base model layers.")
                doc.add_paragraph("- Experimenting with different model architectures or hyperparameters.")
                doc.add_page_break()

                doc.add_paragraph("© Copyright SSO Consultants", style='Body Text')
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Save the document to a BytesIO object
                doc_io = io.BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)

                st.download_button(
                    label="Download Report",
                    data=doc_io,
                    file_name="image_classification_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                st.session_state.report_generated = True
                st.success("Report generated and ready for download!")
    else:
        st.info("Please upload a ZIP file containing your image dataset to begin.")

    st.markdown("---")
    st.markdown("<p style='text-align: center; color: grey;'>© Copyright SSO Consultants</p>", unsafe_allow_html=True)


# --- Authentication Flow ---
if not st.session_state.authenticated:
    login_page()
else:
    main_app()
